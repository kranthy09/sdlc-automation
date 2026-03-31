"""
Document text parser — converts PDF, DOCX, and TXT files into structured
table rows and prose chunks for downstream requirement extraction.

Engines:
  PDF  → pdfplumber — table-aware extraction; prose from non-table regions;
                      OCR fallback via pdf2image + pytesseract for scanned
                      pages (optional — install the ``ocr`` extra to enable)
  DOCX → python-docx — paragraph-level extraction, heading styles preserved
  TXT  → stdlib pathlib

Table extraction (PDF only):
  pdfplumber.find_tables() detects both lattice (bordered) and stream
  (whitespace-aligned) tables per page.  Each table is converted to a list
  of {raw_column_header: cell_value} dicts — one per data row — returned
  in ParseResult.tables.  Prose text is extracted from the remaining page
  area after table regions are excluded via outside_bbox().
  DOCX tables are extracted via doc.tables (top-level tables only).
  TXT always produces tables=[].

OCR fallback:
  Pages whose non-table prose text is shorter than _SCANNED_THRESHOLD chars
  are treated as scanned/image-only pages.  pdf2image + pytesseract are
  attempted; if either dependency (or the underlying poppler/tesseract system
  tools) is unavailable, the page produces no prose chunks and a DEBUG log
  is emitted.  OCR failures are never fatal.

Usage::

    from platform.parsers.docling_parser import DoclingParser

    parser = DoclingParser()
    result = parser.parse(Path("requirements.pdf"))
    # result.tables → list of {col_header: cell_value} dicts (one per row)
    # result.prose  → list[ProseChunk]
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from platform.observability.logger import get_logger
from platform.schemas.errors import ParseError  # noqa: F401 — re-exported

__all__ = [
    "DoclingParser",
    "ParseResult",
    "ProseChunk",
    "ParseError",
]

log = get_logger(__name__)

_MAX_CHUNK_CHARS = 1500
_OVERLAP_SENTENCES = 2
_SCANNED_THRESHOLD = 20  # chars; pages below this after table removal trigger OCR

# Internal item type: (is_heading, text, page_no)
_TextItem = tuple[bool, str, int]


# ---------------------------------------------------------------------------
# Public result types  (API unchanged — callers see same field names)
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ProseChunk:
    """A chunk of prose text with section context.

    Attributes:
        text:        The chunk text (may start with overlap from prior chunk).
        section:     The most recent section heading above this chunk.
        page:        Page number where the chunk starts (1-based).
        char_offset: Cumulative character offset from start of document.
        has_overlap: True when this chunk's leading text repeats the tail of
                     the previous chunk (overlap stitching for retrieval).
    """

    text: str
    section: str
    page: int
    char_offset: int
    has_overlap: bool


@dataclass(frozen=True)
class ParseResult:
    """Output of a successful ``DoclingParser.parse()`` call.

    Attributes:
        tables: Structured table rows extracted from PDF and DOCX files.  Each
                row is a dict mapping raw column header text to cell value.
                Rows with no cell values are dropped.  Always ``[]`` for TXT.
        prose:  Ordered list of prose chunks (≤ 1500 chars each) from
                non-table page regions, with 2-sentence overlap prefix for
                retrieval context continuity.
    """

    tables: list[dict[str, str]]
    prose: list[ProseChunk]


# ---------------------------------------------------------------------------
# DoclingParser
# ---------------------------------------------------------------------------


class DoclingParser:
    """PDF/DOCX/TXT parser backed by pdfplumber and python-docx."""

    def __init__(self) -> None:
        pass

    def parse(self, path: Path) -> ParseResult:
        """Parse a document file into structured table rows and prose chunks.

        PDF parsing uses pdfplumber to:
          1. Extract structured table rows (ParseResult.tables).
          2. Extract prose from non-table page regions (ParseResult.prose).
          3. Attempt OCR on pages that yield no embedded text.

        DOCX extracts both tables (via doc.tables) and prose (via doc.paragraphs).
        TXT produces tables=[] and prose only.

        Args:
            path: Path to a PDF, DOCX, or TXT file.

        Returns:
            ``ParseResult`` with ``tables`` and ``prose`` populated.

        Raises:
            ``ParseError``: If the file cannot be read or parsed.
        """
        try:
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                items, tables = _extract_pdf(path)
            elif suffix == ".docx":
                items, tables = _extract_docx(path)
            else:
                items = _extract_txt(path)
                tables = []
            prose = _chunk(items)
        except ParseError:
            raise
        except Exception as exc:
            raise ParseError(filename=path.name, reason=str(exc)) from exc

        log.debug(
            "parser_parse_ok",
            file=str(path),
            n_chunks=len(prose),
            n_table_rows=len(tables),
        )
        return ParseResult(tables=tables, prose=prose)


# ---------------------------------------------------------------------------
# Format extractors
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path) -> tuple[list[_TextItem], list[dict[str, str]]]:
    """Extract prose items and structured table rows from a PDF using pdfplumber.

    Per-page strategy:
      1. find_tables()     — detect table objects + bounding boxes.
      2. table.extract()   — pull cell values; first row treated as headers.
      3. outside_bbox()    — exclude each table region from the page so prose
                             extraction does not duplicate table cell text.
      4. extract_text()    — pull prose from the remaining (non-table) area.
      5. OCR fallback      — if prose is shorter than _SCANNED_THRESHOLD chars,
                             attempt _ocr_page(); silently skipped if the ocr
                             extra (pdf2image + pytesseract) is not installed.

    Returns:
        (prose_items, table_rows) where prose_items feed _chunk() and
        table_rows are {raw_header: cell_value} dicts ready for
        _map_table_rows_to_canonical() in the ingestion pipeline.
    """
    try:
        import pdfplumber  # noqa: PLC0415
    except ImportError as exc:
        raise ParseError(
            filename=path.name,
            reason="pdfplumber is not installed — run: uv sync --extra ml",
        ) from exc

    items: list[_TextItem] = []
    tables: list[dict[str, str]] = []

    with pdfplumber.open(str(path)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):

            # ------------------------------------------------------------------
            # Step 1 + 2 — table detection and row extraction
            # ------------------------------------------------------------------
            page_table_objs = page.find_tables()

            for table_obj in page_table_objs:
                raw = table_obj.extract()
                # Need at least a header row + one data row
                if not raw or len(raw) < 2:
                    continue

                # First row is column headers; blank cells get placeholder names
                headers = [
                    (str(cell).strip() if cell else f"col_{i}")
                    for i, cell in enumerate(raw[0])
                ]

                for row in raw[1:]:
                    row_dict: dict[str, str] = {}
                    for header, cell in zip(headers, row):
                        val = str(cell).strip() if cell is not None else ""
                        if val:
                            row_dict[header] = val
                    if row_dict:
                        tables.append(row_dict)

            # ------------------------------------------------------------------
            # Step 3 + 4 — prose extraction from non-table regions
            # ------------------------------------------------------------------
            text_page = page
            for table_obj in page_table_objs:
                try:
                    text_page = text_page.outside_bbox(table_obj.bbox)
                except Exception:
                    # outside_bbox can raise on degenerate bboxes (zero area, etc.)
                    pass

            page_text = (text_page.extract_text() or "").strip()

            # ------------------------------------------------------------------
            # Step 5 — OCR fallback for scanned / image-only pages
            # ------------------------------------------------------------------
            if len(page_text) < _SCANNED_THRESHOLD:
                ocr_text = _ocr_page(path, page_no)
                if ocr_text:
                    log.debug(
                        "pdf_ocr_fallback_used", file=str(path), page=page_no
                    )
                    page_text = ocr_text

            if not page_text:
                continue

            for para in page_text.split("\n\n"):
                para = para.strip()
                if para:
                    items.append((False, para, page_no))

    return items, tables


def _ocr_page(path: Path, page_no: int) -> str:
    """OCR a single PDF page using pdf2image + pytesseract.

    Requires the ``ocr`` optional extra (pdf2image, pytesseract) and the
    system packages poppler-utils + tesseract-ocr.

    Returns an empty string — without raising — when:
      - pdf2image or pytesseract are not installed (ImportError)
      - convert_from_path produces no images
      - any other runtime error occurs

    Args:
        path:    Path to the source PDF file.
        page_no: 1-based page number to OCR.
    """
    try:
        from pdf2image import convert_from_path  # noqa: PLC0415
        import pytesseract  # noqa: PLC0415
    except ImportError:
        log.debug("ocr_deps_unavailable", file=str(path), page=page_no)
        return ""

    try:
        images = convert_from_path(str(path), first_page=page_no, last_page=page_no, dpi=300)
        if not images:
            return ""
        return pytesseract.image_to_string(images[0]).strip()
    except Exception as exc:
        log.warning("ocr_page_failed", file=str(path), page=page_no, error=str(exc))
        return ""


def _extract_docx(
    path: Path,
) -> tuple[list[_TextItem], list[dict[str, str]]]:
    """Extract prose items and structured table rows from a DOCX.

    Paragraphs whose style name starts with ``Heading`` or equals ``Title``
    are treated as section headings; everything else is prose.

    ``doc.paragraphs`` returns only body-level paragraphs — table cell
    paragraphs are excluded, so table content never double-counts in prose.
    ``doc.tables`` is iterated separately to extract structured rows:
      - First row → column headers; blank cells → ``col_N`` placeholder.
      - Tables with fewer than 2 rows (header-only) are skipped.
      - Data rows with no non-empty cells are dropped.

    Returns:
        (prose_items, table_rows) where prose_items feed ``_chunk()`` and
        table_rows are ``{raw_header: cell_value}`` dicts ready for
        ``_map_table_rows_to_canonical()`` in the ingestion pipeline.
    """
    import io  # noqa: PLC0415

    from docx import Document  # noqa: PLC0415

    doc = Document(io.BytesIO(path.read_bytes()))
    items: list[_TextItem] = []
    tables: list[dict[str, str]] = []

    # Prose — body-level paragraphs only (table cells excluded by python-docx)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        is_heading = (
            style_name.startswith("Heading") or style_name == "Title"
        )
        items.append((is_heading, text, 1))

    # Tables — top-level tables only (doc.tables does not recurse nested)
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:  # need header + at least 1 data row
            continue
        headers = [
            (rows[0].cells[j].text.strip() or f"col_{j}")
            for j in range(len(rows[0].cells))
        ]
        for row in rows[1:]:
            row_dict: dict[str, str] = {}
            for header, cell in zip(headers, row.cells):
                val = cell.text.strip()
                if val:
                    row_dict[header] = val
            if row_dict:
                tables.append(row_dict)

    return items, tables


def _extract_txt(path: Path) -> list[_TextItem]:
    """Extract text from a plain-text file.

    Splits on double newlines to produce paragraph-sized items.
    """
    text = path.read_text(encoding="utf-8", errors="replace")
    items: list[_TextItem] = []
    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            items.append((False, para, 1))
    return items


# ---------------------------------------------------------------------------
# Chunker  (algorithm identical to the original implementation)
# ---------------------------------------------------------------------------


def _chunk(items: list[_TextItem]) -> list[ProseChunk]:
    """Group text items into ProseChunks with overlap stitching.

    Algorithm:
    1. Walk items in document order.
    2. Track the current section heading (is_heading == True items).
    3. Accumulate prose items into a buffer.
    4. Emit a ProseChunk when the buffer exceeds ``_MAX_CHUNK_CHARS``.
    5. Stitch overlap: last ``_OVERLAP_SENTENCES`` sentences of chunk N are
       prepended to chunk N+1 for retrieval context continuity.
    6. Flush remaining buffer at end of document and at every section change.
    """
    chunks: list[ProseChunk] = []
    section: str = ""
    buf: list[str] = []
    buf_page: int = 1
    char_offset: int = 0
    overlap_prefix: str = ""

    def flush(*, end_section: bool = False) -> None:
        nonlocal buf, overlap_prefix
        if not buf:
            return
        joined = " ".join(buf)
        full = (overlap_prefix + " " + joined).strip() if overlap_prefix else joined
        chunks.append(
            ProseChunk(
                text=full,
                section=section,
                page=buf_page,
                char_offset=char_offset,
                has_overlap=bool(overlap_prefix),
            )
        )
        overlap_prefix = (
            "" if end_section else _last_sentences(joined, _OVERLAP_SENTENCES)
        )
        buf = []

    for is_heading, text, page in items:
        text = text.strip()
        if not text:
            continue

        if is_heading:
            flush(end_section=True)
            section = text
            overlap_prefix = ""
            continue

        if not buf:
            buf_page = page

        buf.append(text)
        char_offset += len(text) + 1

        if sum(len(t) for t in buf) >= _MAX_CHUNK_CHARS:
            flush()

    flush(end_section=True)
    return chunks


def _last_sentences(text: str, n: int) -> str:
    """Return the last *n* sentences from *text* (period-based split)."""
    parts = [s.strip() for s in text.split(". ") if s.strip()]
    tail = parts[-n:] if len(parts) > n else []
    if not tail:
        return ""
    joined = ". ".join(tail)
    return joined if joined.endswith(".") else joined + "."
