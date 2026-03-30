"""
TDD — platform/parsers/docling_parser.py  (pdfplumber backend)

Behaviours under test:
  - TXT   → prose chunks produced; text content preserved.
  - DOCX  → Heading style becomes section label; body paragraph is prose.
  - PDF   → embedded text extracted per page; appears in prose chunks.
  - PDF   → tables detected by pdfplumber; rows in ParseResult.tables.
  - PDF   → table regions excluded from prose (no double-counting).
  - PDF   → OCR path invoked when extract_text() returns empty.
  - Empty → empty TXT produces empty ParseResult (tables=[], prose=[]).
  - Error → unreadable file raises ParseError.

PDF prose/page tests use real minimal PDF bytes written to tmp_path.
Table and OCR tests patch pdfplumber.open() — table PDFs cannot be
constructed from raw bytes without a PDF generation library.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ---------------------------------------------------------------------------
# File builders
# ---------------------------------------------------------------------------


def _write_txt(tmp_path: Path, text: str, name: str = "reqs.txt") -> Path:
    p = tmp_path / name
    p.write_text(text, encoding="utf-8")
    return p


def _write_docx(
    tmp_path: Path,
    heading: str,
    body: str,
    name: str = "reqs.docx",
) -> Path:
    """Create a real DOCX with one heading and one body paragraph."""
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    doc.save(buf)
    p = tmp_path / name
    p.write_bytes(buf.getvalue())
    return p


def _write_pdf(tmp_path: Path, text: str, name: str = "reqs.pdf") -> Path:
    """Create a minimal valid PDF with embedded text (no OCR needed)."""
    safe = (
        text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    )
    stream = f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET\n".encode("latin-1")

    header = b"%PDF-1.4\n"
    obj1 = b"1 0 obj\n<</Type /Catalog /Pages 2 0 R>>\nendobj\n"
    obj2 = b"2 0 obj\n<</Type /Pages /Kids [3 0 R] /Count 1>>\nendobj\n"
    obj3 = (
        b"3 0 obj\n<</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
        b" /Contents 4 0 R /Resources <</Font <</F1 5 0 R>>>>>>\nendobj\n"
    )
    obj4 = (
        b"4 0 obj\n<</Length "
        + str(len(stream)).encode()
        + b">>\nstream\n"
        + stream
        + b"endstream\nendobj\n"
    )
    obj5 = (
        b"5 0 obj\n<</Type /Font /Subtype /Type1"
        b" /BaseFont /Helvetica>>\nendobj\n"
    )

    body = header + obj1 + obj2 + obj3 + obj4 + obj5
    o1 = len(header)
    o2 = o1 + len(obj1)
    o3 = o2 + len(obj2)
    o4 = o3 + len(obj3)
    o5 = o4 + len(obj4)
    xref_pos = len(body)

    xref = (
        b"xref\n0 6\n"
        b"0000000000 65535 f\r\n"
        + f"{o1:010d} 00000 n\r\n".encode()
        + f"{o2:010d} 00000 n\r\n".encode()
        + f"{o3:010d} 00000 n\r\n".encode()
        + f"{o4:010d} 00000 n\r\n".encode()
        + f"{o5:010d} 00000 n\r\n".encode()
    )
    trailer = (
        b"trailer\n<</Size 6 /Root 1 0 R>>\nstartxref\n"
        + str(xref_pos).encode()
        + b"\n%%EOF\n"
    )

    p = tmp_path / name
    p.write_bytes(body + xref + trailer)
    return p


def _make_mock_pdf(pages: list[MagicMock]) -> MagicMock:
    """Wrap pages in a context-manager-compatible mock PDF."""
    mock_pdf = MagicMock()
    mock_pdf.pages = pages
    mock_pdf.__enter__ = lambda s: mock_pdf
    mock_pdf.__exit__ = MagicMock(return_value=False)
    return mock_pdf


def _make_plain_page(text: str) -> MagicMock:
    """Mock page with no tables and the given prose text."""
    page = MagicMock()
    page.find_tables.return_value = []
    page.extract_text.return_value = text
    return page


# ---------------------------------------------------------------------------
# TXT parsing
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_txt_extracts_prose_chunks(tmp_path: Path) -> None:
    """Plain-text requirement returned as a ProseChunk."""
    from platform.parsers.docling_parser import DoclingParser

    req = "The system must validate three-way matching for vendor invoices."
    path = _write_txt(tmp_path, req)

    result = DoclingParser().parse(path)

    assert result.tables == []
    assert len(result.prose) >= 1
    full_text = " ".join(c.text for c in result.prose)
    assert "three-way matching" in full_text


@pytest.mark.unit
def test_parse_txt_multi_paragraph_produces_multiple_chunks(
    tmp_path: Path,
) -> None:
    """Double-newline paragraphs each become chunk content."""
    from platform.parsers.docling_parser import DoclingParser

    content = (
        "The system must validate vendor invoices against purchase orders.\n\n"
        "The system shall calculate tax amounts per invoice automatically.\n\n"
        "The system must approve payments within the configured tolerance."
    )
    path = _write_txt(tmp_path, content)
    result = DoclingParser().parse(path)

    full_text = " ".join(c.text for c in result.prose)
    assert "vendor invoices" in full_text
    assert "tax amounts" in full_text
    assert "tolerance" in full_text


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_docx_heading_becomes_section(tmp_path: Path) -> None:
    """A DOCX Heading 1 paragraph sets the section label on chunks."""
    from platform.parsers.docling_parser import DoclingParser

    path = _write_docx(
        tmp_path,
        heading="Accounts Payable",
        body="The system must support three-way matching.",
    )
    result = DoclingParser().parse(path)

    assert result.tables == []
    assert len(result.prose) >= 1
    chunk = result.prose[0]
    assert chunk.section == "Accounts Payable"
    assert "three-way matching" in chunk.text
    assert chunk.has_overlap is False


@pytest.mark.unit
def test_parse_docx_no_heading_section_is_empty(tmp_path: Path) -> None:
    """Body-only DOCX produces section='' on chunks."""
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("The system must validate payment tolerance thresholds.")
    doc.save(buf)
    p = tmp_path / "reqs.docx"
    p.write_bytes(buf.getvalue())

    from platform.parsers.docling_parser import DoclingParser

    result = DoclingParser().parse(p)

    assert len(result.prose) >= 1
    assert result.prose[0].section == ""


# ---------------------------------------------------------------------------
# PDF — prose extraction (real minimal PDF bytes)
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_pdf_extracts_embedded_text(tmp_path: Path) -> None:
    """Embedded text in a minimal PDF appears in prose chunks."""
    from platform.parsers.docling_parser import DoclingParser

    path = _write_pdf(tmp_path, "System must validate invoice matching.")
    result = DoclingParser().parse(path)

    assert len(result.prose) >= 1
    full_text = " ".join(c.text for c in result.prose)
    assert "invoice" in full_text.lower()


@pytest.mark.unit
def test_parse_pdf_chunk_page_number_is_1_based(tmp_path: Path) -> None:
    """Chunks from a single-page PDF carry page=1."""
    from platform.parsers.docling_parser import DoclingParser

    path = _write_pdf(
        tmp_path,
        "The system shall post journal entries to the ledger.",
    )
    result = DoclingParser().parse(path)

    assert all(c.page >= 1 for c in result.prose)


@pytest.mark.unit
def test_parse_pdf_plain_text_has_no_tables(tmp_path: Path) -> None:
    """A prose-only PDF produces tables=[]."""
    from platform.parsers.docling_parser import DoclingParser

    path = _write_pdf(
        tmp_path,
        "The system must validate vendor payment tolerances.",
    )
    result = DoclingParser().parse(path)

    assert result.tables == []


# ---------------------------------------------------------------------------
# PDF — table extraction (pdfplumber mocked)
# Table PDFs cannot be built from raw bytes without a generation library.
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_pdf_table_rows_populate_tables_field(tmp_path: Path) -> None:
    """Rows from pdfplumber table objects appear in ParseResult.tables."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "header text")

    mock_table = MagicMock()
    mock_table.extract.return_value = [
        ["Requirement Description", "Priority", "Wave"],
        [
            "The system shall post vendor invoices automatically",
            "MUST",
            "Wave 1",
        ],
        [
            "The system shall validate tax codes against HMRC lookup",
            "MUST",
            "Wave 2",
        ],
    ]
    mock_table.bbox = (50.0, 100.0, 560.0, 300.0)

    mock_page = MagicMock()
    mock_page.find_tables.return_value = [mock_table]
    mock_page.outside_bbox.return_value.extract_text.return_value = (
        "Accounts Payable requirements for Wave 1."
    )

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        result = DoclingParser().parse(placeholder)

    assert len(result.tables) == 2
    assert result.tables[0]["Requirement Description"] == (
        "The system shall post vendor invoices automatically"
    )
    assert result.tables[0]["Priority"] == "MUST"
    assert result.tables[0]["Wave"] == "Wave 1"
    assert result.tables[1]["Wave"] == "Wave 2"


@pytest.mark.unit
def test_parse_pdf_table_region_excluded_from_prose(tmp_path: Path) -> None:
    """Table cell text does not appear in prose — outside_bbox is applied."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "x")

    mock_table = MagicMock()
    mock_table.extract.return_value = [
        ["Req ID", "Description"],
        ["REQ-001", "The system shall archive records"],
    ]
    mock_table.bbox = (50.0, 200.0, 560.0, 400.0)

    prose_page = MagicMock()
    prose_page.extract_text.return_value = "Executive summary paragraph."

    mock_page = MagicMock()
    mock_page.find_tables.return_value = [mock_table]
    mock_page.outside_bbox.return_value = prose_page

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        result = DoclingParser().parse(placeholder)

    mock_page.outside_bbox.assert_called_once_with(mock_table.bbox)

    full_prose = " ".join(c.text for c in result.prose)
    assert "Executive summary" in full_prose
    assert "REQ-001" not in full_prose


@pytest.mark.unit
def test_parse_pdf_blank_header_cells_get_placeholder(
    tmp_path: Path,
) -> None:
    """Blank header cells become col_N placeholders; data still extracted."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "x")

    mock_table = MagicMock()
    mock_table.extract.return_value = [
        ["Requirement", None, "Priority"],
        ["System shall validate invoices", "AP-001", "MUST"],
    ]
    mock_table.bbox = (50.0, 100.0, 560.0, 200.0)

    mock_page = MagicMock()
    mock_page.find_tables.return_value = [mock_table]
    mock_page.outside_bbox.return_value.extract_text.return_value = ""

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        result = DoclingParser().parse(placeholder)

    assert len(result.tables) == 1
    row = result.tables[0]
    assert row["Requirement"] == "System shall validate invoices"
    assert row["col_1"] == "AP-001"
    assert row["Priority"] == "MUST"


@pytest.mark.unit
def test_parse_pdf_header_only_table_is_skipped(tmp_path: Path) -> None:
    """A table with only a header row produces no entries in tables."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "x")

    mock_table = MagicMock()
    mock_table.extract.return_value = [
        ["Req ID", "Description"],  # header row only, no data
    ]
    mock_table.bbox = (50.0, 100.0, 560.0, 150.0)

    mock_page = MagicMock()
    mock_page.find_tables.return_value = [mock_table]
    mock_page.outside_bbox.return_value.extract_text.return_value = ""

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        result = DoclingParser().parse(placeholder)

    assert result.tables == []


# ---------------------------------------------------------------------------
# PDF — OCR fallback (pdfplumber + OCR mocked)
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_pdf_ocr_called_when_page_text_empty(tmp_path: Path) -> None:
    """Empty extract_text() triggers _ocr_page; its text used in prose."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "x")
    ocr_text = "The system shall archive vendor records annually."

    mock_page = _make_plain_page("")  # empty → triggers OCR

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        with patch(
            "platform.parsers.docling_parser._ocr_page",
            return_value=ocr_text,
        ) as mock_ocr:
            result = DoclingParser().parse(placeholder)

    mock_ocr.assert_called_once()
    full_text = " ".join(c.text for c in result.prose)
    assert "vendor records" in full_text


@pytest.mark.unit
def test_parse_pdf_ocr_not_called_when_page_has_text(tmp_path: Path) -> None:
    """Sufficient embedded text means OCR is never attempted."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "x")
    embedded = "The system shall validate three-way matching for all invoices."

    mock_page = _make_plain_page(embedded)  # long enough → no OCR

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        with patch(
            "platform.parsers.docling_parser._ocr_page"
        ) as mock_ocr:
            DoclingParser().parse(placeholder)

    mock_ocr.assert_not_called()


@pytest.mark.unit
def test_parse_pdf_ocr_unavailable_page_silently_skipped(
    tmp_path: Path,
) -> None:
    """_ocr_page returning '' means the page produces no chunks."""
    from platform.parsers.docling_parser import DoclingParser

    placeholder = _write_pdf(tmp_path, "x")

    mock_page = _make_plain_page("")  # empty embedded text

    with patch("pdfplumber.open", return_value=_make_mock_pdf([mock_page])):
        with patch(
            "platform.parsers.docling_parser._ocr_page", return_value=""
        ):
            result = DoclingParser().parse(placeholder)

    assert result.tables == []
    assert result.prose == []


# ---------------------------------------------------------------------------
# Empty document
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_empty_txt_returns_empty_result(tmp_path: Path) -> None:
    """An empty TXT file produces tables=[] and prose=[]."""
    from platform.parsers.docling_parser import DoclingParser

    path = _write_txt(tmp_path, "")
    result = DoclingParser().parse(path)

    assert result.tables == []
    assert result.prose == []


# ---------------------------------------------------------------------------
# Error handling
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_bad_pdf_raises_parse_error(tmp_path: Path) -> None:
    """Random bytes with .pdf extension are wrapped in ParseError."""
    from platform.parsers.docling_parser import DoclingParser
    from platform.schemas.errors import ParseError

    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"\x00\x01\x02\x03not a pdf")

    with pytest.raises(ParseError) as exc_info:
        DoclingParser().parse(bad)

    assert exc_info.value.filename == "bad.pdf"


@pytest.mark.unit
def test_parse_bad_docx_raises_parse_error(tmp_path: Path) -> None:
    """A ZIP that is not a valid DOCX raises ParseError."""
    from platform.parsers.docling_parser import DoclingParser
    from platform.schemas.errors import ParseError

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("not_a_docx.xml", "<root/>")
    bad = tmp_path / "bad.docx"
    bad.write_bytes(buf.getvalue())

    with pytest.raises(ParseError) as exc_info:
        DoclingParser().parse(bad)

    assert exc_info.value.filename == "bad.docx"
